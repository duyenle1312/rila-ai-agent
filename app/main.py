from fastapi import FastAPI, File, UploadFile
import mammoth
from docx import Document
from app.services.llm import summarize_text
from app.services.notifier import send_email_notification
from app.services.notion import create_notion_page

app = FastAPI()


def extract_docx_title(file_bytes: bytes) -> str:
    """
    Extract title from DOCX metadata or fallback to first paragraph.
    """
    with open("temp.docx", "wb") as f:
        f.write(file_bytes)

    doc = Document("temp.docx")
    title = doc.core_properties.title or ""
    if not title and doc.paragraphs:
        # first paragraph as fallback
        title = doc.paragraphs[0].text.strip()[:200]
    return title


def convert_docx_to_html(file_bytes: bytes) -> str:
    """
    Convert DOCX content to HTML using Mammoth (preserves headings, bold, italic, lists)
    """
    with open("temp.docx", "wb") as f:
        f.write(file_bytes)

    with open("temp.docx", "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value  # Extracted HTML
        # result.messages contains warnings if any
    return html


@app.get("/")
async def root():
    from app.config import settings
    print("NOTION_API_KEY: ", settings.NOTION_API_KEY)

    return {"message": "Welcome to the Blog Agent API"}


@app.post("/read-docx/")
async def read_docx(file: UploadFile = File(...)):

    # Step 1 — Extract text from DOCX
    if not file.filename.endswith(".docx"):
        return {"error": "Unsupported file type. Only .docx files are supported."}

    file_bytes = await file.read()

    title = extract_docx_title(file_bytes)
    html_content = convert_docx_to_html(file_bytes)

    # Step 2 — Summarize with Gemini
    llm_response = await summarize_text(title=title, html_content=html_content)

    # Step 3 — Send to Notion
    notion_response = await create_notion_page(title=title, slug=llm_response["slug"], ai_summary=llm_response["ai_summary"], seo_keywords=llm_response["seo_keywords"], coverImg=llm_response["cover_imgUrl"], html_content=llm_response["html_content"])

    # Step 4 - Send notification (WhatsApp / Email)
    if notion_response["status"] == "success":
        email_sent = send_email_notification(title, notion_response["url"])
    else:
        print("Failed to create Notion page.")

    return {
        "title": title,
        "status": notion_response["status"],
        "notion_page": notion_response["url"] if notion_response["status"] == "success" else None,
        "filename": file.filename,
        "email_sent": "successfully" if email_sent else "failed"
    }
